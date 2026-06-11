"""File-type readers that produce Markdown.

Each reader takes a ``ScannedFile`` plus a ``ReaderContext`` (which knows
where to copy assets/data files) and returns a ``ReaderResult`` containing
the converted Markdown body and any side-effect notes.

Readers must never raise on a single bad file — they catch their own
exceptions and return a ``ReaderResult`` with ``status='failed'``.
"""

from __future__ import annotations

import csv
import io
import json
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional
from urllib.parse import unquote, urlparse

from . import markdown_utils as mdu
from .scanner import ScannedFile


@dataclass
class ReaderContext:
    """Side-effect targets shared across reader invocations."""

    source_root: Path  # absolute root of the input project
    assets_dir: Path  # where to copy images and other binary referenced assets
    data_dir: Path  # where to copy CSV/XLSX originals


@dataclass
class ReaderResult:
    """The output of converting a single source file."""

    markdown: str
    status: str  # "ok", "skipped", "failed"
    notes: str = ""
    char_count: int = 0
    word_count: int = 0
    copied_assets: List[Path] = field(default_factory=list)


# ----------------------------------------------------------------------------
# Public dispatcher
# ----------------------------------------------------------------------------


def read_file(file: ScannedFile, doc_id: str, ctx: ReaderContext) -> ReaderResult:
    """Convert ``file`` to Markdown. Never raises."""
    try:
        if file.file_type == "text":
            return _read_text(file)
        if file.file_type == "html":
            return _read_html(file, doc_id, ctx)
        if file.file_type == "pdf":
            return _read_pdf(file)
        if file.file_type == "docx":
            return _read_docx(file)
        if file.file_type == "csv":
            return _read_csv(file, doc_id, ctx)
        if file.file_type == "xlsx":
            return _read_xlsx(file, doc_id, ctx)
        if file.file_type == "json":
            return _read_json(file)
        if file.file_type == "image":
            return _read_image(file, doc_id, ctx)
        return ReaderResult(
            markdown="",
            status="skipped",
            notes=f"Unsupported file type: {file.extension}",
        )
    except Exception as exc:  # pragma: no cover - defensive
        tb = traceback.format_exc(limit=3)
        return ReaderResult(
            markdown="",
            status="failed",
            notes=f"Reader crashed: {exc!s} | {tb.splitlines()[-1] if tb else ''}",
        )


# ----------------------------------------------------------------------------
# Text / Markdown
# ----------------------------------------------------------------------------


def _read_text(file: ScannedFile) -> ReaderResult:
    text = _read_text_with_fallback(file.absolute_path)
    text = mdu.normalize_newlines(text)
    return ReaderResult(
        markdown=text,
        status="ok",
        char_count=len(text),
        word_count=len(text.split()),
    )


def _read_text_with_fallback(path: Path) -> str:
    """Read text as UTF-8; on failure, fall back to latin-1 with replacement."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1", errors="replace")
    except OSError as exc:
        raise RuntimeError(f"Could not read text file: {exc}") from exc


# ----------------------------------------------------------------------------
# HTML
# ----------------------------------------------------------------------------


def _read_html(file: ScannedFile, doc_id: str, ctx: ReaderContext) -> ReaderResult:
    from bs4 import BeautifulSoup  # local import to keep startup fast

    raw = _read_text_with_fallback(file.absolute_path)
    # lxml is fastest if available; fall back to html.parser.
    try:
        soup = BeautifulSoup(raw, "lxml")
    except Exception:
        soup = BeautifulSoup(raw, "html.parser")

    for tag_name in ("script", "style", "nav", "footer", "noscript"):
        for tag in soup.find_all(tag_name):
            tag.decompose()

    notes_parts: List[str] = []
    copied_assets: List[Path] = []

    # Rewrite <img src=...> to point at the bundled asset path.
    doc_assets_dir = ctx.assets_dir / doc_id
    for img in soup.find_all("img"):
        original_src = img.get("src", "")
        new_src = _process_html_image(
            original_src=original_src,
            html_file=file.absolute_path,
            doc_assets_dir=doc_assets_dir,
            copied_assets=copied_assets,
            notes_parts=notes_parts,
        )
        if new_src is not None:
            if new_src.startswith("MISSING/"):
                img.replace_with(soup.new_string(f"[Missing image asset: {original_src}]"))
            else:
                img["src"] = new_src

    # Use markdownify to get readable Markdown out of the cleaned soup.
    try:
        from markdownify import markdownify as md_convert  # type: ignore

        body = soup.body if soup.body else soup
        md = md_convert(str(body), heading_style="ATX")
    except Exception:
        # Fallback: extract visible text only.
        md = soup.get_text("\n", strip=True)

    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes="; ".join(notes_parts),
        char_count=len(md),
        word_count=len(md.split()),
        copied_assets=copied_assets,
    )


def _process_html_image(
    original_src: str,
    html_file: Path,
    doc_assets_dir: Path,
    copied_assets: List[Path],
    notes_parts: List[str],
) -> Optional[str]:
    """Try to copy a local image referenced by HTML into the bundle assets dir.

    Returns the new (relative) src to use in the rendered Markdown, or ``None``
    if the original src should be kept (e.g., remote URLs).
    """
    if not original_src:
        return None

    parsed = urlparse(original_src)
    if parsed.scheme in ("http", "https", "data", "mailto"):
        # Leave remote / data URIs untouched.
        return None

    # Try to resolve a local path. Strip any query/fragment.
    candidate_str = unquote(parsed.path) if parsed.path else original_src
    candidate = Path(candidate_str)
    if not candidate.is_absolute():
        candidate = (html_file.parent / candidate).resolve()

    if not candidate.exists() or not candidate.is_file():
        notes_parts.append(f"Missing image asset: {original_src}")
        # Surface the missing-asset note inline too, by using an obvious href.
        return f"MISSING/{original_src}"

    try:
        dest = mdu.unique_destination(doc_assets_dir, candidate.name)
        dest.write_bytes(candidate.read_bytes())
        copied_assets.append(dest)
    except OSError as exc:
        notes_parts.append(f"Failed to copy image {candidate.name}: {exc}")
        return f"MISSING/{original_src}"

    # Use a path relative to the bundle root (assets/DOC_xxxx/file.png).
    return f"assets/{doc_assets_dir.name}/{dest.name}"


# ----------------------------------------------------------------------------
# PDF
# ----------------------------------------------------------------------------


def _read_pdf(file: ScannedFile) -> ReaderResult:
    """Try PyMuPDF, then pypdf, then report extraction failure."""
    md, used = _try_pdf_pymupdf(file.absolute_path)
    if md is None:
        md, used = _try_pdf_pypdf(file.absolute_path)
    if md is None:
        return ReaderResult(
            markdown="[PDF text extraction failed: no PDF backend available. "
            "Install pymupdf or pypdf.]\n",
            status="failed",
            notes="No PDF backend (pymupdf or pypdf) installed.",
        )

    stripped = md.strip()
    notes = f"PDF backend: {used}"
    if len(stripped) < 20:
        md = (
            md
            + "\n\n[PDF text extraction returned little or no text. "
            "OCR may be needed.]\n"
        )
        notes += "; little/no text extracted"

    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes=notes,
        char_count=len(md),
        word_count=len(md.split()),
    )


def _try_pdf_pymupdf(path: Path):
    try:
        import fitz  # PyMuPDF
    except Exception:
        return None, None
    try:
        doc = fitz.open(path)
    except Exception as exc:
        return None, f"pymupdf-open-failed:{exc}"
    parts: List[str] = []
    try:
        for i, page in enumerate(doc, start=1):
            try:
                text = page.get_text("text") or ""
            except Exception:
                text = ""
            parts.append(f"## Page {i}\n\n{text.strip()}\n")
    finally:
        doc.close()
    return "\n".join(parts), "pymupdf"


def _try_pdf_pypdf(path: Path):
    try:
        from pypdf import PdfReader
    except Exception:
        return None, None
    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        return None, f"pypdf-open-failed:{exc}"
    parts: List[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        parts.append(f"## Page {i}\n\n{text.strip()}\n")
    return "\n".join(parts), "pypdf"


# ----------------------------------------------------------------------------
# DOCX
# ----------------------------------------------------------------------------


def _read_docx(file: ScannedFile) -> ReaderResult:
    try:
        import docx  # python-docx
    except ImportError:
        return ReaderResult(
            markdown="[python-docx not installed. Cannot read DOCX.]",
            status="failed",
            notes="python-docx is not installed.",
        )

    document = docx.Document(str(file.absolute_path))
    parts: List[str] = []

    # Walk document body in order so paragraphs and tables interleave correctly.
    body = document.element.body
    paragraphs_by_id = {p._p: p for p in document.paragraphs}
    tables_by_id = {t._tbl: t for t in document.tables}

    for child in body.iterchildren():
        if child in paragraphs_by_id:
            para = paragraphs_by_id[child]
            text = para.text.strip()
            if not text:
                parts.append("")
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                level_digits = "".join(ch for ch in style if ch.isdigit())
                level = int(level_digits) if level_digits else 1
                level = max(1, min(level, 6))
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        elif child in tables_by_id:
            table = tables_by_id[child]
            parts.append(_docx_table_to_markdown(table))

    md = "\n\n".join(p for p in parts if p is not None)
    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        char_count=len(md),
        word_count=len(md.split()),
    )


def _docx_table_to_markdown(table) -> str:
    """Render a python-docx table as a Markdown table."""
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    if not rows:
        return ""
    headers = rows[0]
    body = rows[1:] if len(rows) > 1 else []
    if not body:
        # Single-row table: still render with empty body.
        return mdu.rows_to_markdown_table(headers, [])
    return mdu.rows_to_markdown_table(headers, body)


# ----------------------------------------------------------------------------
# CSV
# ----------------------------------------------------------------------------

PREVIEW_ROW_LIMIT = 25


def _read_csv(file: ScannedFile, doc_id: str, ctx: ReaderContext) -> ReaderResult:
    notes_parts: List[str] = []
    copied: List[Path] = []
    try:
        copied_path = _copy_to_data(file.absolute_path, ctx.data_dir, doc_id)
        copied.append(copied_path)
    except OSError as exc:
        notes_parts.append(f"Failed to copy original CSV: {exc}")

    try:
        import pandas as pd  # type: ignore
    except ImportError:
        return _read_csv_stdlib(file, copied, notes_parts)

    try:
        df = pd.read_csv(file.absolute_path)
    except Exception as exc:
        return _read_csv_stdlib(
            file, copied, notes_parts + [f"pandas read failed: {exc}; using stdlib csv"]
        )

    total_rows = len(df)
    columns = [str(c) for c in df.columns]
    preview = df.head(PREVIEW_ROW_LIMIT).astype(object).where(df.head(PREVIEW_ROW_LIMIT).notna(), "")
    rows = [list(row) for row in preview.itertuples(index=False, name=None)]

    md = _csv_markdown(
        original_filename=file.relative_path.name,
        total_rows=total_rows,
        columns=columns,
        preview_rows=rows,
        copied_path=copied[0] if copied else None,
    )
    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes="; ".join(notes_parts),
        char_count=len(md),
        word_count=len(md.split()),
        copied_assets=copied,
    )


def _read_csv_stdlib(file: ScannedFile, copied: List[Path], notes_parts: List[str]) -> ReaderResult:
    """CSV fallback using the standard library when pandas is unavailable or fails."""
    text = _read_text_with_fallback(file.absolute_path)
    reader = csv.reader(io.StringIO(text))
    all_rows = list(reader)
    if not all_rows:
        md = "[CSV is empty.]\n"
        return ReaderResult(
            markdown=md, status="ok", char_count=len(md), word_count=0, copied_assets=copied
        )
    headers = all_rows[0]
    body = all_rows[1:]
    md = _csv_markdown(
        original_filename=file.relative_path.name,
        total_rows=len(body),
        columns=headers,
        preview_rows=body[:PREVIEW_ROW_LIMIT],
        copied_path=copied[0] if copied else None,
    )
    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes="; ".join(notes_parts),
        char_count=len(md),
        word_count=len(md.split()),
        copied_assets=copied,
    )


def _csv_markdown(
    original_filename: str,
    total_rows: int,
    columns: List[str],
    preview_rows: List,
    copied_path: Optional[Path],
) -> str:
    parts: List[str] = []
    parts.append(f"# CSV: {original_filename}\n")
    parts.append(f"- **Row count:** {total_rows}")
    parts.append(f"- **Column count:** {len(columns)}")
    parts.append(f"- **Columns:** {', '.join(columns) if columns else '(none)'}")
    if copied_path is not None:
        parts.append(f"- **Original copied to:** `data/{copied_path.name}`")
    parts.append("")
    if preview_rows:
        shown = min(len(preview_rows), PREVIEW_ROW_LIMIT)
        parts.append(f"## Preview (first {shown} rows)\n")
        parts.append(mdu.rows_to_markdown_table(columns, preview_rows))
    else:
        parts.append("[No data rows.]")
    return "\n".join(parts)


# ----------------------------------------------------------------------------
# XLSX
# ----------------------------------------------------------------------------


def _read_xlsx(file: ScannedFile, doc_id: str, ctx: ReaderContext) -> ReaderResult:
    notes_parts: List[str] = []
    copied: List[Path] = []
    try:
        copied_path = _copy_to_data(file.absolute_path, ctx.data_dir, doc_id)
        copied.append(copied_path)
    except OSError as exc:
        notes_parts.append(f"Failed to copy original XLSX: {exc}")

    try:
        import pandas as pd  # type: ignore
    except ImportError:
        return ReaderResult(
            markdown="[pandas is not installed; cannot preview XLSX content. "
            "Original file copied to data/.]\n",
            status="ok",
            notes="; ".join(notes_parts + ["pandas missing"]),
            copied_assets=copied,
        )

    try:
        # sheet_name=None → dict of {sheet_name: DataFrame}
        sheets = pd.read_excel(file.absolute_path, sheet_name=None)
    except Exception as exc:
        return ReaderResult(
            markdown=f"[Failed to read XLSX: {exc}]\n",
            status="failed",
            notes="; ".join(notes_parts + [str(exc)]),
            copied_assets=copied,
        )

    parts: List[str] = []
    parts.append(f"# Workbook: {file.relative_path.name}\n")
    if copied:
        parts.append(f"_Original copied to_ `data/{copied[0].name}`\n")

    for sheet_name, df in sheets.items():
        total_rows = len(df)
        columns = [str(c) for c in df.columns]
        head = df.head(PREVIEW_ROW_LIMIT)
        head = head.astype(object).where(head.notna(), "")
        preview_rows = [list(row) for row in head.itertuples(index=False, name=None)]
        parts.append(f"## Sheet: {sheet_name}\n")
        parts.append(f"- **Row count:** {total_rows}")
        parts.append(f"- **Column count:** {len(columns)}")
        parts.append(f"- **Columns:** {', '.join(columns) if columns else '(none)'}")
        parts.append("")
        if preview_rows:
            shown = min(len(preview_rows), PREVIEW_ROW_LIMIT)
            parts.append(f"### Preview (first {shown} rows)\n")
            parts.append(mdu.rows_to_markdown_table(columns, preview_rows))
        else:
            parts.append("[No data rows.]")
        parts.append("")

    md = "\n".join(parts)
    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes="; ".join(notes_parts),
        char_count=len(md),
        word_count=len(md.split()),
        copied_assets=copied,
    )


# ----------------------------------------------------------------------------
# JSON
# ----------------------------------------------------------------------------


def _read_json(file: ScannedFile) -> ReaderResult:
    """Render structured JSON as Markdown with one heading per field.

    Object keys become headings (top level at ``##``, nested objects one
    level deeper), so heading-aware tooling — chunk review, structure
    summaries, and future heading-based chunking rules — sees the record's
    own field boundaries. Values are never dropped: nulls and empty
    containers are rendered explicitly.
    """
    raw = _read_text_with_fallback(file.absolute_path)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        return ReaderResult(
            markdown="",
            status="failed",
            notes=f"Invalid JSON: {exc}",
        )

    parts: List[str] = [f"# JSON: {file.relative_path.name}\n"]
    parts.extend(_json_value_to_markdown(data, level=2))
    md = mdu.normalize_newlines("\n".join(parts))
    return ReaderResult(
        markdown=md,
        status="ok",
        char_count=len(md),
        word_count=len(md.split()),
    )


def _json_value_to_markdown(value: object, level: int) -> List[str]:
    if isinstance(value, dict):
        if not value:
            return ["(empty object)", ""]
        lines: List[str] = []
        for key, item in value.items():
            lines.append(f"{'#' * min(level, 6)} {key}")
            lines.append("")
            lines.extend(_json_value_to_markdown(item, level + 1))
        return lines
    if isinstance(value, list):
        if not value:
            return ["(empty list)", ""]
        if all(not isinstance(item, (dict, list)) for item in value):
            return [f"- {_json_scalar(item)}" for item in value] + [""]
        lines = []
        for position, item in enumerate(value, start=1):
            lines.append(f"{'#' * min(level, 6)} item {position}")
            lines.append("")
            lines.extend(_json_value_to_markdown(item, level + 1))
        return lines
    return [_json_scalar(value), ""]


def _json_scalar(value: object) -> str:
    if value is None:
        return "(null)"
    if isinstance(value, bool):
        return "true" if value else "false"
    return str(value)


# ----------------------------------------------------------------------------
# Images
# ----------------------------------------------------------------------------


def _read_image(file: ScannedFile, doc_id: str, ctx: ReaderContext) -> ReaderResult:
    doc_assets_dir = ctx.assets_dir / doc_id
    try:
        dest = mdu.unique_destination(doc_assets_dir, file.relative_path.name)
        dest.write_bytes(file.absolute_path.read_bytes())
    except OSError as exc:
        return ReaderResult(
            markdown=f"[Image asset copy failed: {exc}]\n",
            status="failed",
            notes=str(exc),
        )

    rel = f"assets/{doc_assets_dir.name}/{dest.name}"
    md = (
        f"# Image: {file.relative_path.name}\n\n"
        f"![{file.relative_path.name}]({rel})\n\n"
        "[Image asset copied. No OCR performed in Version 1.]\n"
    )
    md = mdu.normalize_newlines(md)
    return ReaderResult(
        markdown=md,
        status="ok",
        notes="image copied; no OCR in v1",
        char_count=len(md),
        word_count=len(md.split()),
        copied_assets=[dest],
    )


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------


def _copy_to_data(src: Path, data_dir: Path, doc_id: str) -> Path:
    """Copy ``src`` into ``data_dir`` with a doc-id prefix to avoid collisions."""
    data_dir.mkdir(parents=True, exist_ok=True)
    target_name = f"{doc_id}_{src.name}"
    dest = mdu.unique_destination(data_dir, target_name)
    dest.write_bytes(src.read_bytes())
    return dest
